# Document Parser — supports PDF, TXT, DOCX, Markdown
# Scanned PDFs auto-fallback to MinerU OCR API

from __future__ import annotations

import logging
from dataclasses import dataclass, field
from pathlib import Path

logger = logging.getLogger(__name__)


@dataclass
class Document:
    """Parsed document with metadata."""
    filename: str
    content: str
    page_count: int = 0
    metadata: dict = field(default_factory=dict)


def parse_document(
    filepath: str | Path,
    max_pages: int | None = None,
    page_range: tuple[int, int] | None = None,
) -> Document:
    """Parse a document file, autodetecting format by extension.

    Supports: .txt, .md, .pdf, .docx

    Args:
        filepath: Path to the document file.
        max_pages: For PDFs, only parse the first N pages (None = all).
        page_range: For image PDFs, OCR only pages [start, end] (1-indexed).
                   Used by knowledge processing to OCR a single chapter.
    """
    filepath = Path(filepath)
    suffix = filepath.suffix.lower()

    if suffix == ".txt":
        return _parse_txt(filepath)
    elif suffix == ".md":
        return _parse_txt(filepath)
    elif suffix == ".pdf":
        return _parse_pdf(filepath, max_pages=max_pages, page_range=page_range)
    elif suffix == ".docx":
        return _parse_docx(filepath)
    else:
        raise ValueError(f"Unsupported file format: {suffix}")


def _parse_txt(filepath: Path) -> Document:
    content = filepath.read_text(encoding="utf-8", errors="replace")
    lines = content.splitlines()
    return Document(
        filename=filepath.name,
        content=content,
        page_count=len(lines) // 40 or 1,
        metadata={"format": "text", "line_count": len(lines)},
    )


def _parse_pdf(
    filepath: Path,
    max_pages: int | None = None,
    page_range: tuple[int, int] | None = None,
) -> Document:
    """Parse PDF using pymupdf. Multi-phase fallback for scanned PDFs."""
    # Phase 1: try embedded text extraction (fast, works for digital PDFs)
    doc = _parse_pdf_pymupdf_raw(filepath, max_pages)
    if doc.content.strip():
        return doc

    # ---- Scanned PDF: no embedded text ----
    body_doc: Document | None = None

    if page_range:
        _ocr_pages = page_range
    elif max_pages:
        _ocr_pages = (1, min(max_pages, 20))
    else:
        _ocr_pages = (1, 20)

    # Phase 2a: MinerU cloud OCR (fast, best quality, GPU-accelerated).
    # Phase 2b: EasyOCR local (fallback, ~8s/page, no Paddle dependency).
    # Phase 2c: Tesseract local (last resort).

    # Try MinerU first
    logger.info("PDF has no embedded text layer, trying MinerU cloud OCR pages %d-%d",
                 _ocr_pages[0], _ocr_pages[1])
    mineru_text = _parse_pdf_ocr(filepath, page_range=_ocr_pages)
    if mineru_text and mineru_text.content.strip():
        body_doc = mineru_text
        logger.info("[parser] MinerU OCR succeeded: %d chars", len(mineru_text.content))
    else:
        # Fallback to EasyOCR
        logger.info("MinerU failed/empty, falling back to EasyOCR")
        body_doc = _parse_pdf_unstructured(filepath, _ocr_pages[0], _ocr_pages[1])

    # Phase 2d: VLM TOC scan → chapter title list (try regardless of OCR result)
    logger.info("Scanning for table of contents with VLM (%s pages)",
                 str(max_pages) if max_pages else "all")
    toc_doc = _parse_pdf_vlm(filepath, max_pages)

    # Merge: VLM chapter list + page ranges → metadata, body text → content.
    if toc_doc and toc_doc.metadata.get("vlm_chapters"):
        vlm_list = toc_doc.metadata["vlm_chapters"]
        vlm_ranges = toc_doc.metadata.get("vlm_chapter_ranges", [])
        if body_doc and body_doc.content.strip():
            body_doc.metadata["vlm_chapters"] = vlm_list
            body_doc.metadata["vlm_chapter_ranges"] = vlm_ranges
            body_doc.metadata["vlm_body_start"] = toc_doc.metadata.get("vlm_body_start", 0)
            body_doc.metadata["parser"] = "local_ocr+vlm_toc"
            logger.info("[parser] VLM %d chapters (with%s page ranges) + OCR %d chars",
                         len(vlm_list),
                         "" if vlm_ranges else "out",
                         len(body_doc.content))
            return body_doc
        else:
            logger.info("[parser] VLM %d chapters (with%s page ranges), no body text",
                         len(vlm_list),
                         "" if vlm_ranges else "out")
            return toc_doc

    # VLM failed: body text only
    if body_doc and body_doc.content.strip():
        return body_doc

    # Phase 3: VLM + local OCR both failed → MinerU cloud (last resort)
    logger.info("Local OCR + VLM both failed, trying MinerU OCR (%s pages)",
                 str(max_pages) if max_pages else "all")
    ocr_doc = _parse_pdf_ocr(filepath, max_pages)
    if ocr_doc and ocr_doc.content.strip():
        return ocr_doc

    # All approaches failed
    logger.warning("All PDF parsing phases failed for %s", filepath.name)
    return doc


def _parse_pdf_vlm(filepath: Path, max_pages: int | None = None) -> Document | None:
    """Scan first N pages with VLM to find and read the table of contents.

    The VLM's ONLY job is to locate the TOC within the page images and list
    every chapter title it finds. Body text extraction is handled by local
    OCR (Phase 2). This separates concerns: VLM does what it's good at
    (visual pattern recognition), local OCR does what it's good at
    (bulk text extraction).

    Returns a Document whose content is a clean chapter title list,
    or None if no TOC/chapters found.
    """
    import base64
    import io as _io
    import time as _time

    try:
        import fitz
    except ImportError:
        return None

    try:
        doc = fitz.open(str(filepath))
        total = len(doc)
        limit = min(max_pages or total, total, 20)

        # Single call: 20 pages at 150 DPI ≈ 28K tokens, fits under 38K limit
        all_images = []
        for i in range(limit):
            page = doc[i]
            pix = page.get_pixmap(dpi=150)
            buf = _io.BytesIO(pix.tobytes("jpeg"))
            img_b64 = base64.b64encode(buf.read()).decode("utf-8")
            all_images.append(img_b64)

        doc.close()

        if not all_images:
            return None

        image_parts = [
            {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{img}"}}
            for img in all_images
        ]

        from src.config import Configuration
        from langchain_openai import ChatOpenAI
        from langchain_core.messages import HumanMessage

        config = Configuration.from_env()
        vision_model = config.llm_vision_model or config.llm_model_id
        vision_url = config.llm_vision_base_url or config.llm_base_url
        vision_key = config.llm_vision_api_key or config.llm_api_key

        if not config.llm_vision_model:
            logger.warning(
                "[VLM] llm_vision_model not configured, falling back to %s",
                vision_model,
            )

        llm = ChatOpenAI(
            model=vision_model,
            api_key=vision_key,
            base_url=vision_url,
            temperature=0.0,
            max_tokens=2000,  # just chapter titles, not full text
        )

        # Prompt: find TOC, list chapter titles WITH page numbers.
        # The page numbers let us calculate chapter ranges so knowledge
        # processing can OCR only the selected chapter's pages.
        #
        # IMPORTANT: the VLM also identifies which image number (1-indexed,
        # cover = image 1) the first chapter's actual body content starts on.
        # This offset corrects for the gap between printed page numbers and
        # PDF pages (caused by cover, copyright, preface, TOC front matter).
        prompt = (
            f"这是《{filepath.name}》前{limit}页的扫描图片（图片1=封面/第1页，图片2=第2页……以此类推）。"
            "你的任务：\n\n"
            "1. 在图片中找到目录页，列出所有一级章节标题及起始页码。\n"
            "   目录页特征：通常有'目录'或'Contents'标题，"
            "下方列出各章节名及对应页码。\n"
            "   格式：章节标题 | 页码\n"
            "   例如：第一章 静电场的基本规律 | 1\n\n"
            "2. 判断第一个章节的**正文内容**从第几张图片开始。\n"
            "   （不是目录中列出的页码，而是图片编号。目录页不是正文，\n"
            "   封面/版权页/前言/目录 这些都算前置页。）\n"
            "   格式：BODY_START: 数字\n"
            "   例如：BODY_START: 9   （表示第9张图片是第一章正文的第一页）\n\n"
            "只输出以上内容，不要其他。如果找不到目录，输出'未找到'。"
        )

        _t0 = _time.time()
        msg = HumanMessage(content=[{"type": "text", "text": prompt}] + image_parts)
        resp = llm.invoke([msg])
        resp_text = resp.content if hasattr(resp, "content") else str(resp)

        logger.info(
            "[VLM] TOC scan: %d pages, %.1fs, %d chars: %s",
            limit, _time.time() - _t0, len(resp_text),
            resp_text[:300].replace("\n", " | "),
        )

        if not resp_text or "未找到" in resp_text or len(resp_text.strip()) < 10:
            logger.warning("[VLM] no TOC/chapters found in %d pages", limit)
            return None

        # Parse BODY_START offset from VLM response.
        # BODY_START tells us which PDF page (1-indexed) the first chapter's
        # actual body content begins on — after cover, copyright, preface, TOC.
        import re as _re
        body_start = 0
        bs_match = _re.search(r'BODY_START\s*:\s*(\d+)', resp_text, _re.IGNORECASE)
        if bs_match:
            body_start = int(bs_match.group(1))
            logger.info("[VLM] detected BODY_START = %d (PDF page where Ch1 body begins)", body_start)

        # Parse chapter titles and page numbers from VLM response.
        # Expected format per line: "第一章 静电场的基本规律 | 1"
        chapter_ranges: list[dict] = []
        for raw_line in resp_text.strip().splitlines():
            line = raw_line.strip().lstrip("-•·1234567890.、 ")
            if not line or len(line) < 4:
                continue
            # Split by last "|" or any page-number-like suffix
            parts = line.rsplit("|", 1)
            title = parts[0].strip() if parts else line
            page_str = parts[1].strip() if len(parts) > 1 else ""
            # Extract first number from page_str (handle "1-2", "1  ", etc.)
            page_match = _re.search(r'(\d+)', page_str)
            page_num = int(page_match.group(1)) if page_match else 0
            if title and page_num > 0:
                chapter_ranges.append({
                    "title": title,
                    "start_page": page_num,
                })

        if not chapter_ranges:
            # Fallback: no page numbers — just titles
            lines = []
            for raw_line in resp_text.strip().splitlines():
                title = raw_line.strip().lstrip("-•·1234567890.、 ")
                if title and len(title) >= 4:
                    lines.append(title)
            if not lines:
                return None
            chapter_ranges = [{"title": t, "start_page": 0} for t in lines]

        # Calculate end pages: Chapter N ends at (Chapter N+1 start - 1).
        # Last chapter ends at total page count.
        chapter_ranges.sort(key=lambda c: c["start_page"])
        for i, ch in enumerate(chapter_ranges):
            if i + 1 < len(chapter_ranges):
                ch["end_page"] = chapter_ranges[i + 1]["start_page"] - 1
            else:
                ch["end_page"] = total

        titles = [c["title"] for c in chapter_ranges]
        logger.info(
            "[VLM] parsed %d chapters with page ranges: %s",
            len(chapter_ranges),
            ", ".join(f"{c['title'][:12]}... p{c['start_page']}-{c['end_page']}"
                      for c in chapter_ranges[:5]),
        )

        # Format as JSON for chapterizer compatibility
        import json
        chapters_data = {
            "chapters": [
                {"title": c["title"], "start_marker": c["title"]}
                for c in chapter_ranges
            ]
        }
        content = json.dumps(chapters_data, ensure_ascii=False)
        return Document(
            filename=filepath.name,
            content=content,
            page_count=total,
            metadata={
                "format": "pdf",
                "parser": "vlm_toc",
                "pages_scanned": limit,
                "chapters_found": len(chapter_ranges),
                "vlm_chapters": titles,
                "vlm_chapter_ranges": chapter_ranges,
                "vlm_body_start": body_start,  # PDF page where Ch1 body begins (1-indexed)
            },
        )

    except Exception as e:
        logger.warning("VLM TOC scan failed: %s", e)
        return None


def _vlm_enhance_ocr(
    ocr_result: list,
    page_img,
    doc_name: str = "",
    page_num: int = 0,
    confidence_threshold: float = 0.5,
) -> list:
    """Enhance EasyOCR results by re-recognizing low-confidence regions with Qwen-VL.

    EasyOCR confidence drops on formulas, special symbols, and degraded CJK.
    This function crops those regions from the page image, sends them to a
    vision model (Qwen-VL) which handles formulas naturally, and replaces
    the low-confidence text with the VLM output.

    Returns the modified ocr_result list (same format).
    """
    from src.config import Configuration
    from langchain_openai import ChatOpenAI
    from langchain_core.messages import HumanMessage
    import base64, io as _io

    # Find low-confidence regions
    low_conf = [(i, r) for i, r in enumerate(ocr_result) if r[2] < confidence_threshold]
    if not low_conf:
        return ocr_result

    # Init VLM (lazy, once per function call)
    config = Configuration.from_env()
    vision_model = config.llm_vision_model or config.llm_model_id
    vision_url = config.llm_vision_base_url or config.llm_base_url
    vision_key = config.llm_vision_api_key or config.llm_api_key

    if not config.llm_vision_model:
        logger.debug("[vlm-enhance] no vision model configured, skipping")
        return ocr_result

    llm = ChatOpenAI(
        model=vision_model,
        api_key=vision_key,
        base_url=vision_url,
        temperature=0.0,
        max_tokens=2000,  # batch multi-region output needs more tokens
    )

    # Batch all low-confidence crops into a single VLM call
    import cv2
    crops_b64: list[str] = []
    crop_indices: list[int] = []
    for idx, (bbox, text, conf) in low_conf:
        try:
            x1, y1 = int(bbox[0][0]), int(bbox[0][1])
            x2, y2 = int(bbox[2][0]), int(bbox[2][1])
            h, w = page_img.shape[:2]
            x1, y1 = max(0, x1 - 5), max(0, y1 - 5)
            x2, y2 = min(w, x2 + 5), min(h, y2 + 5)
            crop = page_img[y1:y2, x1:x2]
            if crop.size == 0:
                continue
            _, buf = cv2.imencode(".png", crop)
            crops_b64.append(base64.b64encode(buf.tobytes()).decode("utf-8"))
            crop_indices.append(idx)
        except Exception:
            continue

    if not crops_b64:
        return ocr_result

    # Build multi-image prompt
    image_parts = []
    region_labels = []
    for i in range(len(crops_b64)):
        image_parts.append({
            "type": "image_url",
            "image_url": {"url": f"data:image/png;base64,{crops_b64[i]}"},
        })
        region_labels.append(f"[{i}]")

    prompt = (
        f"这是同一页内{len(crops_b64)}个低质量OCR区域的裁图。\n"
        "请按编号逐个识别每个区域，输出格式：\n"
        "[0] 识别结果\n[1] 识别结果\n...\n"
        "如果包含数学公式，用LaTeX（行内$...$，块级$$...$$）。\n"
        "只输出编号和内容，不要加解释。"
    )

    msg = HumanMessage(content=[{"type": "text", "text": prompt}] + image_parts)
    try:
        resp = llm.invoke([msg])
        vlm_output = resp.content.strip() if hasattr(resp, "content") else str(resp).strip()

        # Parse batched response: [N] text
        import re as _re
        parsed = {}
        pattern = _re.compile(r'\[(\d+)\]\s*(.*?)(?=\[\d+\]|\Z)', _re.DOTALL)
        for m in pattern.finditer(vlm_output):
            idx_str = m.group(1)
            text = m.group(2).strip()
            if idx_str.isdigit() and text:
                parsed[int(idx_str)] = text

        enhanced_count = 0
        for batch_i, orig_idx in enumerate(crop_indices):
            vlm_text = parsed.get(batch_i, "").strip()
            if vlm_text:
                ocr_result[orig_idx] = (ocr_result[orig_idx][0], vlm_text, 0.7)
                enhanced_count += 1

        if enhanced_count > 0:
            logger.info("[vlm-enhance] %s p%d: enhanced %d/%d low-confidence regions (batched)",
                         doc_name, page_num, enhanced_count, len(low_conf))

    except Exception as e:
        logger.warning("[vlm-enhance] batch VLM failed: %s", e)

    if enhanced_count > 0:
        logger.info("[vlm-enhance] %s p%d: enhanced %d/%d low-confidence regions",
                     doc_name, page_num, enhanced_count, len(low_conf))

    return ocr_result


def _parse_pdf_unstructured(
    filepath: Path,
    start_page: int = 1,
    end_page: int = 20,
) -> Document:
    """Parse scanned PDF pages with EasyOCR + VLM enhancement. 3-page concurrent."""
    try:
        import fitz
    except ImportError:
        logger.warning("[easyocr] pymupdf not available")
        return Document(filepath.name, "", 0, {"parser": "easyocr_error"})

    try:
        import easyocr, numpy as np, cv2
    except ImportError:
        logger.warning("[easyocr] easyocr not installed")
        return Document(filepath.name, "", 0, {"parser": "easyocr_error"})

    try:
        import time as _time, concurrent.futures
        _t0 = _time.time()
        doc = fitz.open(str(filepath))
        total = len(doc)
        _start = max(0, start_page - 1)
        _end = min(end_page, total)
        page_count = _end - _start
        OCVLM_WORKERS = 3

        reader = easyocr.Reader(["ch_sim"], gpu=False)

        def _process_page(i: int):
            try:
                pix = doc[i].get_pixmap(dpi=200)
                img = np.frombuffer(pix.tobytes("png"), dtype=np.uint8)
                img = cv2.imdecode(img, cv2.IMREAD_COLOR)
                result = reader.readtext(img)
                try:
                    result = _vlm_enhance_ocr(result, img, filepath.name, i + 1)
                except Exception:
                    pass
                text = "".join(r[1] for r in result)
                return (i, text.strip() if text.strip() else None)
            except Exception as e:
                logger.warning("[easyocr] page %d failed: %s", i + 1, e)
                return (i, None)

        logger.info("[easyocr] %d pages, %d workers", page_count, OCVLM_WORKERS)
        results: dict[int, str] = {}
        pages_done = 0

        with concurrent.futures.ThreadPoolExecutor(max_workers=OCVLM_WORKERS) as ex:
            futures = {ex.submit(_process_page, i): i for i in range(_start, _end)}
            for future in concurrent.futures.as_completed(futures):
                i, text = future.result()
                pages_done += 1
                if text:
                    results[i] = text
                if pages_done % 5 == 0:
                    logger.info("[easyocr] %d/%d pages in %.1fs",
                                 pages_done, page_count, _time.time() - _t0)

        doc.close()
        elapsed = _time.time() - _t0
        all_text = [results[i] for i in sorted(results) if results[i]]
        content = "\n\n".join(all_text)
        logger.info("[easyocr] pages %d-%d: %d/%d pages, %.1fs, %d chars",
                     start_page, end_page, pages_done, page_count, elapsed, len(content))

        return Document(
            filename=filepath.name, content=content, page_count=total,
            metadata={"format": "pdf", "parser": "easyocr", "ocr_engine": "easyocr_crnn",
                      "dpi": 200, "elapsed_s": round(elapsed, 1)},
        )

    except Exception as e:
        logger.error("[easyocr] OCR failed: %s", e)
        return Document(filepath.name, "", 0, {"parser": "easyocr_error"})

def _parse_pdf_pymupdf_raw(filepath: Path, max_pages: int | None = None) -> Document:
    """Extract embedded text via pymupdf (fast, no OCR)."""
    try:
        import fitz
    except ImportError:
        logger.warning("pymupdf not available, falling back to PyPDF2")
        return _parse_pdf_pypdf2(filepath, max_pages)

    doc = fitz.open(str(filepath))
    total = len(doc)
    limit = min(max_pages, total) if max_pages else total
    pages: list[str] = []

    for page in doc[:limit]:
        # get_text("text") preserves reading order better than PyPDF2
        text = page.get_text("text")
        if text and text.strip():
            pages.append(text.strip())

    doc.close()

    content = "\n\n".join(pages)
    return Document(
        filename=filepath.name,
        content=content,
        page_count=total,
        metadata={"format": "pdf", "pages": total, "pages_parsed": len(pages), "parser": "pymupdf"},
    )


def _parse_pdf_toc(filepath: Path, max_pages: int | None = None) -> Document | None:
    """Extract the PDF's built-in table of contents (outline) via pymupdf.

    Many textbooks embed a TOC as PDF outline metadata — this is instant
    and requires no OCR, making it ideal for chapter detection when the
    document body is scanned images.

    Returns a Document whose content is a formatted chapter list,
    or None if the PDF has no outline.
    """
    try:
        import fitz
    except ImportError:
        return None

    try:
        doc = fitz.open(str(filepath))
        toc = doc.get_toc(simple=True)  # list of [level, title, page]
        total_pages = len(doc)
        doc.close()

        if not toc:
            return None

        # Filter to first-level entries within page limit
        max_page = max_pages or float("inf")
        seen = set()
        chapters = []
        for level, title, page in toc:
            title = title.strip()
            if not title:
                continue
            if page > max_page:
                continue
            # Only take first-level (level=1) chapters; skip duplicates
            if level == 1 and title not in seen:
                seen.add(title)
                chapters.append(f"第{len(chapters)+1}章 {title} (第{page}页)")

        if not chapters:
            # No level-1 entries in range — include all levels as flat list
            for level, title, page in toc:
                title = title.strip()
                if title and page <= max_page and title not in seen:
                    seen.add(title)
                    indent = "  " * (level - 1) if level > 1 else ""
                    chapters.append(f"{indent}{title} (第{page}页)")

        if not chapters:
            return None

        content = "【PDF 内置目录】\n\n" + "\n".join(chapters)
        # Store raw TOC entries for direct chapter creation (bypass LLM)
        raw_toc = [
            {"title": ch_title, "page": ch_page, "level": ch_level}
            for ch_level, ch_title, ch_page in toc
            if ch_title.strip() and ch_page <= max_page
        ]
        return Document(
            filename=filepath.name,
            content=content,
            page_count=total_pages,
            metadata={
                "format": "pdf",
                "parser": "pymupdf_toc",
                "toc_entries": len(chapters),
                "total_toc": len(toc),
                "raw_toc": raw_toc,
            },
        )
    except Exception as e:
        logger.warning("PDF TOC extraction failed: %s", e)
        return None


def _parse_pdf_pypdf2(filepath: Path, max_pages: int | None = None) -> Document:
    """Legacy PyPDF2 parser kept as fallback."""
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        raise ImportError("PyPDF2 not installed. Run: pip install PyPDF2")
    
    reader = PdfReader(str(filepath))
    total = len(reader.pages)
    limit = min(max_pages, total) if max_pages else total
    pages: list[str] = []
    for page in reader.pages[:limit]:
        text = page.extract_text()
        if text:
            pages.append(text)
    
    content = "\n\n".join(pages)
    return Document(
        filename=filepath.name,
        content=content,
        page_count=total,
        metadata={"format": "pdf", "pages": total, "pages_parsed": len(pages), "parser": "pypdf2"},
    )


def _parse_docx(filepath: Path) -> Document:
    try:
        from docx import Document as DocxDocument
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")
    
    doc = DocxDocument(str(filepath))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    content = "\n\n".join(paragraphs)
    return Document(
        filename=filepath.name,
        content=content,
        page_count=len(paragraphs) // 30 or 1,
        metadata={"format": "docx", "paragraph_count": len(paragraphs)},
    )

def _parse_pdf_ocr(filepath: Path, max_pages: int | None = None, page_range: tuple[int, int] | None = None) -> Document:
    """Parse scanned/image-based PDF using MinerU cloud OCR.

    Calls MinerU API (requires MINERU_API_TOKEN in .env).
    Falls back gracefully if token is not configured or API fails.
    """
    import fitz

    # Try MinerU cloud API first
    try:
        from src.tools.ocr import parse_pdf as mineru_parse
        text = mineru_parse(str(filepath), max_pages=max_pages, page_range=page_range)
        if text.strip():
            pdf_doc = fitz.open(str(filepath))
            total = len(pdf_doc)
            pdf_doc.close()
            logger.info("MinerU OCR: extracted %d chars", len(text))
            return Document(
                filename=filepath.name,
                content=text,
                page_count=total,
                metadata={"format": "pdf", "pages": total, "parser": "mineru"},
            )
        logger.warning("MinerU returned empty text")
    except ImportError:
        logger.info("MinerU OCR tool not available")
    except Exception as e:
        logger.warning("MinerU OCR failed: %s", e)

    # All OCR failed
    pdf_doc = fitz.open(str(filepath))
    total = len(pdf_doc)
    pdf_doc.close()
    return Document(
        filename=filepath.name,
        content="",
        page_count=total,
        metadata={"format": "pdf", "pages": total, "parser": "none"},
    )
